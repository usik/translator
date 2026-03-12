import asyncio
import io

from docx import Document
from docx.shared import Pt

from ..exceptions import OutputError


class DocxConverter:
    format_name = "docx"
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    file_extension = ".docx"

    def _generate_sync(self, text: str, **options) -> bytes:
        font_family = options.get("font_family", "NanumGothic")
        font_size = options.get("font_size", 12)

        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = font_family
        font.size = Pt(font_size)

        for paragraph_text in text.split("\n"):
            doc.add_paragraph(paragraph_text)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    async def convert(self, text: str, **options) -> bytes:
        try:
            return await asyncio.to_thread(self._generate_sync, text, **options)
        except Exception as e:
            raise OutputError(f"DOCX generation failed: {e}") from e

    def _convert_structured_sync(self, source_file: bytes, segments: list) -> bytes:
        doc = Document(io.BytesIO(source_file))
        seg_map = {s.id: s.text for s in segments}

        for i, para in enumerate(doc.paragraphs):
            seg_id = f"p{i}"
            if seg_id in seg_map:
                if para.runs:
                    para.runs[0].text = seg_map[seg_id]
                    for run in para.runs[1:]:
                        run.text = ""

        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    seg_id = f"t{ti}_r{ri}_c{ci}"
                    if seg_id in seg_map:
                        for para in cell.paragraphs:
                            if para.runs:
                                para.runs[0].text = seg_map[seg_id]
                                for run in para.runs[1:]:
                                    run.text = ""
                                break

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    async def convert_structured(self, source_file: bytes, segments: list) -> bytes:
        try:
            return await asyncio.to_thread(
                self._convert_structured_sync, source_file, segments,
            )
        except Exception as e:
            raise OutputError(f"DOCX structured conversion failed: {e}") from e
