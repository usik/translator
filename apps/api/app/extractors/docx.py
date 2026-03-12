import asyncio
import base64
import io

from docx import Document

from ..schemas import ExtractionResult, TextSegment


class BasicDocxExtractor:
    """Fallback DOCX extractor using python-docx (no API key required)."""

    supported_extensions = [".docx"]
    supported_content_types = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ]

    def _extract_sync(self, file_bytes: bytes) -> str:
        doc = Document(io.BytesIO(file_bytes))
        parts: list[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)

        return "\n".join(parts)

    async def extract(self, file_bytes: bytes, filename: str) -> ExtractionResult:
        text = await asyncio.to_thread(self._extract_sync, file_bytes)

        return ExtractionResult(
            text=text,
            source_filename=filename,
            source_format="docx",
            metadata={},
        )

    def _extract_structured_sync(self, file_bytes: bytes) -> tuple[list[TextSegment], str]:
        doc = Document(io.BytesIO(file_bytes))
        segments: list[TextSegment] = []
        text_parts: list[str] = []

        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                segments.append(TextSegment(
                    id=f"p{i}",
                    text=para.text,
                    path=f"paragraphs[{i}]",
                ))
                text_parts.append(para.text)

        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    if cell.text.strip():
                        segments.append(TextSegment(
                            id=f"t{ti}_r{ri}_c{ci}",
                            text=cell.text,
                            path=f"tables[{ti}].rows[{ri}].cells[{ci}]",
                        ))
                        text_parts.append(cell.text)

        return segments, "\n".join(text_parts)

    async def extract_structured(self, file_bytes: bytes, filename: str) -> ExtractionResult:
        segments, text = await asyncio.to_thread(self._extract_structured_sync, file_bytes)
        return ExtractionResult(
            text=text,
            source_filename=filename,
            source_format="docx",
            segments=segments,
            source_file_b64=base64.b64encode(file_bytes).decode("ascii"),
        )


class DocxExtractor:
    """DOCX extractor using Mistral OCR."""

    supported_extensions = [".docx"]
    supported_content_types = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ]

    def __init__(self, mistral_api_key: str) -> None:
        from mistralai import Mistral
        self._client = Mistral(api_key=mistral_api_key)

    def _extract_sync(self, file_bytes: bytes) -> object:
        encoded = base64.standard_b64encode(file_bytes).decode("utf-8")
        document_url = f"data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{encoded}"

        return self._client.ocr.process(
            model="mistral-ocr-latest",
            document={
                "type": "document_url",
                "document_url": document_url,
            },
        )

    async def extract(self, file_bytes: bytes, filename: str) -> ExtractionResult:
        ocr_response = await asyncio.to_thread(self._extract_sync, file_bytes)

        pages = ocr_response.pages
        text_parts = [page.markdown for page in pages]
        full_text = "\n\n".join(text_parts)

        return ExtractionResult(
            text=full_text,
            source_filename=filename,
            source_format="docx",
            page_count=len(pages),
            metadata={},
        )
